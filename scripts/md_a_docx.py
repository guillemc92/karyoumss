# -*- coding: utf-8 -*-
"""Convierte los entregables Markdown de este repositorio a .docx.

No pretende ser un conversor general: cubre exactamente lo que usan los
documentos de entrega —titulos, parrafos, listas, tablas, bloques de codigo,
negrita, cursiva y `codigo en linea`— y nada mas. Es deliberado: un conversor
completo seria una dependencia externa mas (pandoc) que hay que instalar en
cualquier maquina donde se regenere la entrega.

    python scripts/md_a_docx.py docs/M7_ACTIVIDAD2_SUITE_MEDIDA.md

Escribe el .docx junto al .md, con el mismo nombre.
"""
import io
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

GRIS = RGBColor(0x44, 0x44, 0x44)
MONO = 'Consolas'

# `texto` en linea, **negrita**, *cursiva*. El orden importa: la negrita antes
# que la cursiva, porque ** contiene *.
TROZOS = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)')


def escribir_inline(parrafo, texto):
    """Reparte el texto en tramos con o sin formato."""
    for trozo in TROZOS.split(texto):
        if not trozo:
            continue
        if trozo.startswith('`') and trozo.endswith('`'):
            r = parrafo.add_run(trozo[1:-1])
            r.font.name = MONO
            r.font.size = Pt(9.5)
        elif trozo.startswith('**') and trozo.endswith('**'):
            parrafo.add_run(trozo[2:-2]).bold = True
        elif trozo.startswith('*') and trozo.endswith('*'):
            parrafo.add_run(trozo[1:-1]).italic = True
        else:
            # Los enlaces se quedan con su texto: un .docx que se imprime no
            # gana nada con la URL, y la ruta relativa no resuelve fuera del repo.
            parrafo.add_run(re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', trozo))


def fila_de(linea):
    """Celdas de una fila de tabla markdown, sin las barras de los extremos."""
    return [c.strip() for c in linea.strip().strip('|').split('|')]


def es_separador(linea):
    return bool(re.match(r'^\|[\s:|-]+\|$', linea.strip()))


def bloque_codigo(doc, lineas):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('\n'.join(lineas))
    r.font.name = MONO
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS


def tabla(doc, filas):
    if not filas:
        return
    t = doc.add_table(rows=len(filas), cols=len(filas[0]))
    t.style = 'Light Grid Accent 1'
    for i, fila in enumerate(filas):
        for j, celda in enumerate(fila[:len(filas[0])]):
            p = t.cell(i, j).paragraphs[0]
            escribir_inline(p, celda)
            if i == 0:
                for r in p.runs:
                    r.bold = True
    doc.add_paragraph()


def convertir(ruta_md, ruta_docx):
    lineas = io.open(ruta_md, encoding='utf-8').read().split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    i, filas_tabla = 0, []
    while i < len(lineas):
        linea = lineas[i]
        despojada = linea.strip()

        # Una tabla termina en cuanto aparece algo que no es una fila. Se vuelca
        # ANTES de escribir esa linea, o el parrafo siguiente saldria delante.
        if filas_tabla and not despojada.startswith('|'):
            pendientes, filas_tabla = filas_tabla, []
            tabla(doc, pendientes)

        if despojada.startswith('```'):                      # bloque de codigo
            i += 1
            cuerpo = []
            while i < len(lineas) and not lineas[i].strip().startswith('```'):
                cuerpo.append(lineas[i])
                i += 1
            bloque_codigo(doc, cuerpo)
        elif despojada.startswith('|'):                      # tabla
            if not es_separador(despojada):
                filas_tabla.append(fila_de(despojada))
        elif despojada.startswith('#'):                      # titulo
            nivel = len(despojada) - len(despojada.lstrip('#'))
            # Un titulo no lleva formato dentro: se limpian las marcas en vez
            # de dejarlas visibles («Unit — `rag_qa.py`»).
            titulo = despojada[nivel:].strip().replace('**', '').replace('`', '')
            doc.add_heading(titulo, level=min(nivel, 4))
        elif despojada in ('---', '***'):                    # separador
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('· · ·').font.color.rgb = GRIS
        elif re.match(r'^[-*] ', despojada):                 # lista
            escribir_inline(doc.add_paragraph(style='List Bullet'), despojada[2:])
        elif re.match(r'^\d+\. ', despojada):
            escribir_inline(doc.add_paragraph(style='List Number'),
                            re.sub(r'^\d+\. ', '', despojada))
        elif despojada:                                      # parrafo
            escribir_inline(doc.add_paragraph(), despojada)

        i += 1

    if filas_tabla:
        tabla(doc, filas_tabla)

    doc.save(ruta_docx)
    return ruta_docx


if __name__ == '__main__':
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    entrada = sys.argv[1]
    salida = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(entrada)[0] + '.docx'
    print('escrito %s (%d bytes)' % (convertir(entrada, salida),
                                     os.path.getsize(salida)))
