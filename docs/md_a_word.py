"""Convierte los Markdown de entrega a .docx.

    ../backend-clinic/.venv/Scripts/python md_a_word.py ENTREGA_TOOL_CALLING.md

Cubre lo que usan los documentos de entrega: encabezados, párrafos, negrita e
`inline code`, tablas, bloques de código y citas. No pretende ser un conversor
general — pretende que el documento se regenere solo cuando cambie la fuente,
en vez de mantenerse a mano en dos formatos.
"""
from __future__ import annotations

import difflib
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

CODIGO_FONDO = RGBColor(0x1F, 0x2A, 0x37)
ROJO = RGBColor(0xC0, 0x00, 0x00)
NEGRITA_O_CODE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')


def lineas_cambiadas(origen: Path, ref: str) -> set[int]:
    """Indices (0-based) de las lineas nuevas o modificadas respecto a `ref`.

    Sirve para pintar en rojo solo lo que cambio, y asi poder trasladar las
    correcciones a mano a una maqueta externa sin releer el documento entero.
    """
    ruta_git = subprocess.run(
        ['git', 'ls-files', '--full-name', origen.name],
        cwd=origen.parent, capture_output=True, text=True,
    ).stdout.strip()
    if not ruta_git:
        return set()
    previo = subprocess.run(
        ['git', 'show', f'{ref}:{ruta_git}'],
        cwd=origen.parent, capture_output=True, text=True, encoding='utf-8',
    )
    if previo.returncode != 0:
        return set()

    antes = previo.stdout.splitlines()
    ahora = origen.read_text(encoding='utf-8').splitlines()
    cambiadas: set[int] = set()
    matcher = difflib.SequenceMatcher(None, antes, ahora, autojunk=False)
    for etiqueta, _, _, j1, j2 in matcher.get_opcodes():
        if etiqueta in ('replace', 'insert'):
            cambiadas.update(range(j1, j2))
    return cambiadas


def _texto_con_formato(parrafo, texto: str, rojo: bool = False) -> None:
    """Escribe `texto` aplicando **negrita** e `inline code`.

    `rojo=True` marca el fragmento como modificado respecto a la version previa.
    """
    def _pinta(run):
        if rojo:
            run.font.color.rgb = ROJO
        return run

    for trozo in NEGRITA_O_CODE.split(texto):
        if not trozo:
            continue
        if trozo.startswith('**') and trozo.endswith('**'):
            _pinta(parrafo.add_run(trozo[2:-2])).bold = True
        elif trozo.startswith('`') and trozo.endswith('`'):
            run = _pinta(parrafo.add_run(trozo[1:-1]))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            _pinta(parrafo.add_run(trozo))


def _tabla(doc: Document, filas: list[str], rojas: list[bool] | None = None) -> None:
    celdas = [[c.strip() for c in f.strip().strip('|').split('|')] for f in filas]
    rojas = rojas or [False] * len(filas)
    cabecera, cuerpo = celdas[0], celdas[2:]      # celdas[1] es el separador ---
    rojo_cuerpo = rojas[2:]
    tabla = doc.add_table(rows=1, cols=len(cabecera))
    tabla.style = 'Light Grid Accent 1'
    for i, titulo in enumerate(cabecera):
        celda = tabla.rows[0].cells[i]
        celda.text = ''
        _texto_con_formato(celda.paragraphs[0], titulo, rojas[0])
        for run in celda.paragraphs[0].runs:
            run.bold = True
    for n, fila in enumerate(cuerpo):
        nuevas = tabla.add_row().cells
        marca = rojo_cuerpo[n] if n < len(rojo_cuerpo) else False
        for i, valor in enumerate(fila[:len(cabecera)]):
            nuevas[i].text = ''
            _texto_con_formato(nuevas[i].paragraphs[0], valor, marca)


def _bloque_codigo(doc: Document, lineas: list[str], rojo: bool = False) -> None:
    parrafo = doc.add_paragraph()
    parrafo.paragraph_format.left_indent = Pt(14)
    parrafo.paragraph_format.space_after = Pt(10)
    run = parrafo.add_run('\n'.join(lineas))
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = ROJO if rojo else CODIGO_FONDO


def convertir(origen: Path, destino: Path, marcadas: set[int] | None = None) -> None:
    marcadas = marcadas or set()
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    lineas = origen.read_text(encoding='utf-8').splitlines()

    # Las líneas de texto seguidas forman UN párrafo, como en Markdown. Hacerlo
    # así no es cosmético: la negrita puede abrir en una línea y cerrar en la
    # siguiente (`**texto\nmas texto**`), y tratando cada línea por separado los
    # asteriscos salían literales en el documento.
    buffer: list[tuple[int, str]] = []

    def _volcar():
        if not buffer:
            return
        texto = ' '.join(t for _, t in buffer)
        rojo_parrafo = any(n in marcadas for n, _ in buffer)
        _texto_con_formato(doc.add_paragraph(), texto, rojo_parrafo)
        buffer.clear()

    i = 0
    while i < len(lineas):
        linea = lineas[i]
        rojo = i in marcadas

        if linea.startswith('```'):                       # bloque de código
            _volcar()
            cierre = i + 1
            while cierre < len(lineas) and not lineas[cierre].startswith('```'):
                cierre += 1
            # El bloque entero se marca si cambió cualquiera de sus líneas: un
            # bloque de código medio rojo se lee peor que uno entero.
            _bloque_codigo(doc, lineas[i + 1:cierre],
                           any(n in marcadas for n in range(i, cierre + 1)))
            i = cierre + 1
            continue

        if linea.startswith('|'):                          # tabla
            _volcar()
            fin = i
            while fin < len(lineas) and lineas[fin].startswith('|'):
                fin += 1
            _tabla(doc, lineas[i:fin], [n in marcadas for n in range(i, fin)])
            doc.add_paragraph()
            i = fin
            continue

        if linea.startswith('#'):                          # encabezado
            _volcar()
            nivel = len(linea) - len(linea.lstrip('#'))
            enc = doc.add_heading(linea.lstrip('# ').strip(), level=min(nivel, 4))
            if rojo:
                for run in enc.runs:
                    run.font.color.rgb = ROJO
        elif linea.strip() == '---':
            _volcar()
            doc.add_paragraph()
        elif linea.startswith('> '):                       # cita
            _volcar()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _texto_con_formato(p, linea[2:], rojo)
            for run in p.runs:
                run.italic = True
        elif linea.startswith(('- ', '* ')):               # viñeta
            _volcar()
            _texto_con_formato(doc.add_paragraph(style='List Bullet'), linea[2:], rojo)
        elif re.match(r'^\d+\. ', linea):                  # numerada
            _volcar()
            _texto_con_formato(doc.add_paragraph(style='List Number'),
                               re.sub(r'^\d+\. ', '', linea), rojo)
        elif linea.strip():
            buffer.append((i, linea))                      # se acumula
        else:
            _volcar()                                      # línea en blanco

        i += 1

    _volcar()
    doc.save(destino)
    marca = f'  ({len(marcadas)} lineas en rojo)' if marcadas else ''
    print(f'  {destino.name}  <-  {origen.name}{marca}')


if __name__ == '__main__':
    base = Path(__file__).parent
    args = sys.argv[1:]

    # --marcar[=REF]: pinta en ROJO lo que cambio respecto a REF (por defecto
    # HEAD, es decir, lo que todavia no esta commiteado). Pensado para trasladar
    # las correcciones a una maqueta externa sin releer el documento entero.
    ref = None
    resto = []
    for a in args:
        if a.startswith('--marcar'):
            ref = a.split('=', 1)[1] if '=' in a else 'HEAD'
        else:
            resto.append(a)

    fuentes = resto or ['ENTREGA_TOOL_CALLING.md',
                        'ENTREGA_TOOL_CALLING_UNA_PAGINA.md']
    print('Generando Word:' + (f'  (marcando cambios vs {ref})' if ref else ''))
    for nombre in fuentes:
        origen = base / nombre
        marcadas = lineas_cambiadas(origen, ref) if ref else set()
        convertir(origen, origen.with_suffix('.docx'), marcadas)
